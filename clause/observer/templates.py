"""Fillers for the four ARS Stage-1 templates.

Each filler takes:
    blank_path: Path to the blank .docx template
    fields: dict from observer.extract.extract_fields()
    auditor: dict of inputs the QMS can't provide (audit dates, team, etc.)
    out_path: where to write the filled .docx
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.table import _Cell


# --- shared helpers ------------------------------------------------------

def _append_to_cell(cell: _Cell, value: str) -> None:
    """Append `value` as a new paragraph in the cell, copying first-run formatting."""
    if not value:
        return
    if value in cell.text:
        return
    para = cell.add_paragraph(value)
    if cell.paragraphs and cell.paragraphs[0].runs:
        src = cell.paragraphs[0].runs[0]
        for run in para.runs:
            run.font.name = src.font.name
            if src.font.size:
                run.font.size = src.font.size


def _set_cell_text(cell: _Cell, value: str) -> None:
    """Replace the entire cell content with `value`, preserving first-run formatting."""
    if not cell.paragraphs:
        cell.add_paragraph(value)
        return
    p = cell.paragraphs[0]
    # Drop extra paragraphs
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)
    if not p.runs:
        p.add_run(value)
        return
    p.runs[0].text = value
    for r in p.runs[1:]:
        r.text = ""


def _replace_in_paragraph(para, old: str, new: str) -> bool:
    """Substitute `old` with `new` inside a paragraph, preserving first-run formatting."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if not para.runs:
        para.add_run(new_full)
        return True
    para.runs[0].text = new_full
    for r in para.runs[1:]:
        r.text = ""
    return True


def _fill_label_cells(doc, label_to_value: dict[str, str]) -> None:
    """Walk all tables; for each cell whose stripped text matches a label key,
    append the corresponding value (first occurrence wins)."""
    filled: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            seen_in_row: set[str] = set()
            for cell in row.cells:
                key = cell._tc.getroottree().getpath(cell._tc)
                if key in seen_in_row:
                    continue
                seen_in_row.add(key)
                text = cell.text.strip()
                if text in label_to_value and text not in filled:
                    _append_to_cell(cell, label_to_value[text])
                    filled.add(text)


# --- F-006 Audit Plan (Stage-1 + Stage-2 share most labels) --------------

# Each label key maps to a field name; we accept both Stage-1 and Stage-2 variants.
_PLAN_LABELS = {
    "organization_name": ["Organization Name", "Organization Name:"],
    "postal_address": ["Postal Address", "Postal Address:"],
    "audit_site": [
        "Audit Site (and Full detail of Temporary Sites)",
        "Audit Site (and Full detail of Temporary Sites if applicable)",
    ],
    "audit_scope": ["Audit Scope"],
    "contact_person": ["Contact Person:"],
    "contact_number": ["Contact No.:"],
}


def fill_audit_plan(blank: Path, fields: dict, auditor: dict, out: Path) -> None:
    doc = Document(str(blank))
    label_map: dict[str, str] = {}
    for field, variants in _PLAN_LABELS.items():
        value = fields.get(field, "")
        for v in variants:
            label_map[v] = value
    _fill_label_cells(doc, label_map)
    # Fill auditor-supplied bits in the same table
    audit_team = auditor.get("audit_team", [])
    audit_dates = auditor.get("audit_dates", [])  # list of date strings, one per slot
    man_days = auditor.get("audit_man_days", "")

    # Audit Date / Audit Standard cell (right side of header)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt.startswith("Audit Date:") and "Audit Standard:" in txt:
                    if auditor.get("audit_date_range"):
                        _append_to_cell(cell, f"Date(s): {auditor['audit_date_range']}")
                    break
                if txt.startswith("Audit Man-days:"):
                    if man_days and "_" in txt:
                        _set_cell_text(cell, f"Audit Man-days: {man_days}")

    # Audit Team Member rows: row 5 has labels (Lead Auditor / Auditor / Tech Expert),
    # row 6 has empty cells where names go.
    if audit_team:
        team_table = doc.tables[0]
        if len(team_table.rows) >= 7:
            row6 = team_table.rows[6]
            # Pick the first 4 distinct cells to insert names into
            distinct_cells: list[_Cell] = []
            seen_keys: set[str] = set()
            for cell in row6.cells:
                k = cell._tc.getroottree().getpath(cell._tc)
                if k in seen_keys:
                    continue
                seen_keys.add(k)
                distinct_cells.append(cell)
            for cell, member in zip(distinct_cells[1:], audit_team):
                # member is dict {name, role}
                _set_cell_text(cell, member.get("name", ""))

    # Audit Date column in the schedule table: rows 10..15, col 0 (currently 0800/0830/...)
    if audit_dates:
        plan_table = doc.tables[0]
        for row_idx, date_str in zip(range(10, 16), audit_dates):
            if row_idx < len(plan_table.rows):
                first_cell = plan_table.rows[row_idx].cells[0]
                _set_cell_text(first_cell, date_str)

    doc.save(str(out))


# --- F-007(a) Intimation Letter ------------------------------------------

def fill_intimation(blank: Path, fields: dict, auditor: dict, out: Path) -> None:
    doc = Document(str(blank))

    org = fields.get("organization_name", "")
    contact = fields.get("contact_person", "")
    designation = fields.get("contact_designation", "") or "—"
    address = fields.get("postal_address", "").replace("\n", ", ")
    contract_no = auditor.get("contract_number", "TBD")
    schedule_date = auditor.get("audit_date_range", "TBD")
    intimation_date = auditor.get("intimation_date", "")

    for para in doc.paragraphs:
        text = para.text
        if "Intimation Date:" in text and intimation_date:
            _replace_in_paragraph(para, "Intimation Date:", f"Intimation Date: {intimation_date}")
        elif text.startswith("Mr."):
            _set_paragraph_text(para, f"Mr./Ms. {contact}, Designation: {designation}")
        elif text.startswith("M/s."):
            _set_paragraph_text(para, f"M/s. {org}")
        elif text.startswith("Address-"):
            _set_paragraph_text(para, f"Address- {address}")
        elif "vide contract no." in text:
            _set_paragraph_text(
                para,
                f"We would like to inform you with reference to vide contract no. "
                f"{contract_no} that the above audit has been scheduled on "
                f"{schedule_date} as agreed with you.",
            )

    # Audit team table (5 rows: header + 4 role rows: Lead Auditor / Auditor / Expert / Observer)
    audit_team = auditor.get("audit_team", [])
    if audit_team and doc.tables:
        team_table = doc.tables[0]
        # Group team members by normalized role.
        by_role: dict[str, list[str]] = {}
        for m in audit_team:
            role = (m.get("role") or "").strip().lower()
            by_role.setdefault(role, []).append(m.get("name", ""))
        for row in team_table.rows[1:]:  # skip header
            cells = row.cells
            if len(cells) < 3:
                continue
            row_role = cells[2].text.strip().lower()
            # Strict match: exact role label, then fall back to single-word match
            # ("expert" -> "technical expert").
            candidates = (
                by_role.get(row_role)
                or by_role.get(row_role.replace("technical ", ""))
                or next(
                    (
                        names
                        for r, names in by_role.items()
                        if r and r == row_role
                    ),
                    None,
                )
            )
            if candidates:
                _set_cell_text(cells[1], candidates.pop(0))

    doc.save(str(out))


def _set_paragraph_text(para, value: str) -> None:
    """Overwrite paragraph text into run 0, clearing other runs."""
    if not para.runs:
        para.add_run(value)
        return
    para.runs[0].text = value
    for r in para.runs[1:]:
        r.text = ""


# --- F-008 Attendance Sheet ----------------------------------------------

def fill_attendance(blank: Path, fields: dict, auditor: dict, out: Path) -> None:
    doc = Document(str(blank))

    audit_date = auditor.get("audit_date_range", "")
    client_name = fields.get("organization_name", "")
    client_ref = auditor.get("client_reference", "")
    audit_stage = auditor.get("audit_stage", "Stage 1")  # Stage 1 / Stage 2 / Surveillance

    for para in doc.paragraphs:
        text = para.text
        if "Date:-" in text and audit_date:
            _replace_in_paragraph(para, "Date:-", f"Date:- {audit_date}")
        elif "Client Name:" in text and client_name:
            _replace_in_paragraph(para, "Client Name:", f"Client Name: {client_name}")
        elif "Client Reference:" in text:
            # Preserve tab structure: "Client Reference:\tARS/<ref>\t\t\tAudit: <stage>"
            _replace_in_paragraph(para, "ARS/", f"ARS/{client_ref}")
            _replace_in_paragraph(
                para,
                "Stage 1 / Stage 2 / Surveillance",
                audit_stage,
            )

    # Pre-populate audit team rows in the attendance table
    audit_team = auditor.get("audit_team", [])
    if audit_team and doc.tables:
        att = doc.tables[0]
        for member, row in zip(audit_team, att.rows[1:]):
            cells = row.cells
            if len(cells) >= 3:
                _set_cell_text(cells[1], member.get("name", ""))
                _set_cell_text(cells[2], member.get("role", ""))
                _set_cell_text(cells[3], "Audit Team")

    doc.save(str(out))


# --- F-011 / F-012 IMS Audit Report (Stage-1 + Stage-2 share verification table)

REPORT_VERIFICATION_MAP = {
    # cell label in column 0 -> field name (from extract or auditor)
    "Client ref number": ("auditor", "client_reference"),
    "Name": ("fields", "organization_name"),
    "Address": ("fields", "postal_address"),
    "Employee": ("fields", "employee_count"),
    "Contact Person Name": ("fields", "contact_person"),
    "Contact number": ("fields", "contact_number"),
    "E-mail ID": ("fields", "contact_email"),
    "Scope": ("fields", "audit_scope"),
    "IAF Code/NACE": ("fields", "iaf_code"),
    "Audit Man-days": ("auditor", "audit_man_days"),
    "Audit Date": ("auditor", "audit_date_range"),
}


def fill_audit_report(blank: Path, fields: dict, auditor: dict, out: Path) -> None:
    doc = Document(str(blank))

    # Table 0 is the Verification table (subject / information / verification).
    if not doc.tables:
        doc.save(str(out))
        return
    verif_table = doc.tables[0]
    for row in verif_table.rows:
        if len(row.cells) < 2:
            continue
        subject = row.cells[0].text.strip()
        if subject in REPORT_VERIFICATION_MAP:
            source, key = REPORT_VERIFICATION_MAP[subject]
            value = (fields if source == "fields" else auditor).get(key, "")
            if value:
                _set_cell_text(row.cells[1], str(value))
        elif subject == "Audit Team":
            members = auditor.get("audit_team", [])
            if members:
                team_str = "; ".join(
                    f"{m.get('name','')} ({m.get('role','')})" for m in members
                )
                _set_cell_text(row.cells[1], team_str)

    doc.save(str(out))
